# -*- coding: utf-8 -*-
import os
import sys
import re
import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import win32com.client

def parse_markdown(filepath):
    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # フロントマターを除去
    front_matter_pattern = re.compile(r'^---\s*\n(.*?)\n---\s*\n', re.DOTALL)
    match = front_matter_pattern.match(content)
    if match:
        body = content[match.end():]
    else:
        body = content
        
    lines = body.split('\n')
    parsed = []
    
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
            
        # 見出し判定
        if stripped.startswith('# '):
            parsed.append(('h1', stripped[2:]))
        elif stripped.startswith('## '):
            parsed.append(('h2', stripped[3:]))
        elif stripped.startswith('### '):
            parsed.append(('h3', stripped[4:]))
        # 箇条書き判定
        elif stripped.startswith('* ') or stripped.startswith('- '):
            parsed.append(('bullet', stripped[2:]))
        # 引用 (警告・注意ブロック等)
        elif stripped.startswith('> '):
            clean_text = stripped[2:]
            if clean_text.startswith('[!'):
                # 警告タグ
                tag_match = re.match(r'^\[!(.*?)\]\s*(.*)', clean_text)
                if tag_match:
                    parsed.append(('callout', (tag_match.group(1), tag_match.group(2))))
                else:
                    parsed.append(('callout', ('NOTE', clean_text)))
            else:
                parsed.append(('quote', clean_text))
        # 通常の段落
        else:
            parsed.append(('paragraph', line))
            
    return parsed

def add_paragraph_with_inline_formatting(doc_or_cell, text, style='Normal', alignment=None):
    # 空の段落を定義したスタイルで追加
    p = doc_or_cell.add_paragraph(style=style)
    if alignment is not None:
        p.alignment = alignment
        
    # ** で囲まれた文字列を太字（bold=True）にする
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            bold_text = part[2:-2]
            run = p.add_run(bold_text)
            run.bold = True
        else:
            p.add_run(part)
    return p

def build_docx(parsed_data, output_path):
    doc = docx.Document()
    
    # 余白設定 (上下左右 25.4mm = 1インチ)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # スタイルのカスタム
    # Normalスタイル（本文）
    style_normal = doc.styles['Normal']
    font_normal = style_normal.font
    font_normal.name = 'Yu Gothic'
    font_normal.size = Pt(10.5)
    font_normal.color.rgb = RGBColor(0x1F, 0x29, 0x33) # 濃いグレー
    
    # 見出し1スタイル
    style_h1 = doc.styles['Heading 1']
    font_h1 = style_h1.font
    font_h1.name = 'Yu Gothic'
    font_h1.size = Pt(20)
    font_h1.bold = True
    font_h1.color.rgb = RGBColor(0x0F, 0x76, 0x6E) # 濃いエメラルドグリーン
    
    # 見出し2スタイル
    style_h2 = doc.styles['Heading 2']
    font_h2 = style_h2.font
    font_h2.name = 'Yu Gothic'
    font_h2.size = Pt(14)
    font_h2.bold = True
    font_h2.color.rgb = RGBColor(0x0F, 0x76, 0x6E)

    # 見出し3スタイル
    style_h3 = doc.styles['Heading 3']
    font_h3 = style_h3.font
    font_h3.name = 'Yu Gothic'
    font_h3.size = Pt(12)
    font_h3.bold = True
    font_h3.color.rgb = RGBColor(0x11, 0x5E, 0x59)

    for item_type, content in parsed_data:
        if item_type == 'h1':
            p = add_paragraph_with_inline_formatting(doc, content, style='Heading 1', alignment=WD_ALIGN_PARAGRAPH.CENTER)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(18)
        elif item_type == 'h2':
            p = add_paragraph_with_inline_formatting(doc, content, style='Heading 2')
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
        elif item_type == 'h3':
            p = add_paragraph_with_inline_formatting(doc, content, style='Heading 3')
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
        elif item_type == 'bullet':
            p = add_paragraph_with_inline_formatting(doc, content, style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
        elif item_type == 'callout':
            tag, text = content
            table = doc.add_table(rows=1, cols=1)
            table.autofit = False
            table.columns[0].width = Inches(6.5)
            cell = table.cell(0, 0)
            
            # 背景色設定
            from docx.oxml import parse_xml
            from docx.oxml.ns import nsdecls
            fill_color = "FFF1F0" if tag == 'IMPORTANT' else "F8FAFC"
            shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
            cell._tc.get_or_add_tcPr().append(shading_elm)
            
            # 枠線 (左側だけ太線)
            border_color = "B42318" if tag == 'IMPORTANT' else "0F766E"
            borders_elm = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                f'<w:left w:val="single" w:sz="24" w:space="0" w:color="{border_color}"/>'
                f'<w:top w:val="none"/>'
                f'<w:right w:val="none"/>'
                f'<w:bottom w:val="none"/>'
                f'</w:tcBorders>'
            )
            cell._tc.get_or_add_tcPr().append(borders_elm)
            
            cp = cell.paragraphs[0]
            run_tag = cp.add_run(f"【{tag}】 ")
            run_tag.bold = True
            if tag == 'IMPORTANT':
                run_tag.font.color.rgb = RGBColor(0xB4, 0x23, 0x18)
            else:
                run_tag.font.color.rgb = RGBColor(0x0F, 0x76, 0x6E)
            
            # コールアウト内のテキストに対しても太字パースを実行
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    bold_text = part[2:-2]
                    run = cp.add_run(bold_text)
                    run.bold = True
                else:
                    cp.add_run(part)
            
            cp.paragraph_format.space_after = Pt(0)
            
            # スペース空け
            doc.add_paragraph().paragraph_format.space_after = Pt(6)
        elif item_type == 'quote':
            p = add_paragraph_with_inline_formatting(doc, content, style='Normal')
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.space_after = Pt(6)
        else:
            p = add_paragraph_with_inline_formatting(doc, content, style='Normal')
            p.paragraph_format.space_after = Pt(6)
            
    doc.save(output_path)
    print(f"Successfully generated DOCX at {output_path}")

def convert_docx_to_pdf(docx_path, pdf_path):
    # Microsoft Word をバックグラウンドで起動して PDF エクスポート
    word = win32com.client.DispatchEx("Word.Application")
    word.Visible = False
    try:
        abs_docx = os.path.abspath(docx_path)
        abs_pdf = os.path.abspath(pdf_path)
        
        doc = word.Documents.Open(abs_docx)
        # 17 = PDF format
        doc.SaveAs(abs_pdf, FileFormat=17)
        doc.Close()
        print(f"Successfully converted {docx_path} to {pdf_path}")
    except Exception as e:
        print(f"Error during PDF conversion: {e}")
        raise e
    finally:
        word.Quit()

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python convert_to_office.py <markdown_file>")
        sys.exit(1)
        
    md_file = sys.argv[1]
    if not os.path.exists(md_file):
        print(f"Error: {md_file} does not exist.")
        sys.exit(1)
        
    base_name = os.path.splitext(os.path.basename(md_file))[0]
    
    # 出力先ディレクトリの確保
    base_dir = r"C:\Users\yert1\Documents\agy\10_Medical\Patient-information"
    docx_dir = os.path.join(base_dir, "docx")
    pdf_dir = os.path.join(base_dir, "PDF")
    
    if not os.path.exists(docx_dir):
        os.makedirs(docx_dir)
    if not os.path.exists(pdf_dir):
        os.makedirs(pdf_dir)
        
    docx_path = os.path.join(docx_dir, f"{base_name}.docx")
    pdf_path = os.path.join(pdf_dir, f"{base_name}.pdf")
    
    parsed = parse_markdown(md_file)
    build_docx(parsed, docx_path)
    convert_docx_to_pdf(docx_path, pdf_path)
